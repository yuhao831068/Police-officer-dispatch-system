import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches
from database import DatabaseConnection
from utils import get_team_order, format_date, get_rank_restrictions


class ShiftManager:
    """警察局排班管理系統"""

    def __init__(self):
        """初始化排班管理器"""
        self.db = DatabaseConnection()
        self.shift_start_date = datetime(2024, 1, 6).date()
        self.shift_patterns = {
            '123檔期': 0,
            '456檔期': 14,
            '789檔期': 7
        }

    def connect(self):
        """連接資料庫"""
        self.db.connect()

    def disconnect(self):
        """關閉資料庫連接"""
        self.db.disconnect()

    def view_daily_shifts(self, specific_date):
        """
        查看某日的所有班表

        Args:
            specific_date: 指定日期

        Returns:
            DataFrame: 包含該日所有班別資訊的DataFrame
        """
        query = """
        SELECT 
            CASE shift_name
                WHEN 'A班' THEN '01_A班'
                WHEN 'B班' THEN '02_B班'
                WHEN 'C班' THEN '03_C班'
                WHEN 'D班' THEN '04_D班'
                WHEN 'E班' THEN '05_E班'
                WHEN '上值日' THEN '06_上值日'
                WHEN '下值日' THEN '07_下值日'
                WHEN '日值日官' THEN '08_日值日官'
                WHEN '夜值日官' THEN '09_夜值日官'
                WHEN '值班副大隊長' THEN '10_值班副大隊長'
                WHEN '日勤務管理員' THEN '11_日勤務管理員'
                WHEN '夜勤務管理員' THEN '12_夜勤務管理員'
                WHEN '日械彈管理員' THEN '13_日械彈管理員'
                WHEN '夜械彈管理員' THEN '14_夜械彈管理員'
                ELSE shift_name
            END as sort_order,
            shift_name,
            e.S_ID, 
            e.name, 
            e.job_rank,
            s.team_order,
            s.day_order,
            e.current_shift
        FROM Shift s
        JOIN Employee_Shift e ON s.S_ID = e.S_ID
        WHERE s.shift_date = %s
        ORDER BY sort_order, s.team_order, s.day_order
        """
        try:
            # 使用傳統方式獲取數據
            self.db.get_cursor().execute(query, (specific_date,))
            rows = self.db.get_cursor().fetchall()

            # 如果沒有數據，返回空的DataFrame
            if not rows:
                return pd.DataFrame(columns=[
                    'shift_name', 'S_ID', 'name', 'job_rank',
                    'team_order', 'day_order', 'current_shift'
                ])

            # 將結果轉換為DataFrame
            df = pd.DataFrame(rows, columns=[
                'sort_order', 'shift_name', 'S_ID', 'name', 'job_rank',
                'team_order', 'day_order', 'current_shift'
            ])

            # 移除排序用的欄位
            df = df.drop('sort_order', axis=1)

            # 設定顯示選項
            pd.set_option('display.max_rows', None)
            pd.set_option('display.max_columns', None)
            pd.set_option('display.width', None)
            pd.set_option('display.max_colwidth', None)

            return df

        except Exception as err:
            print(f"查詢錯誤: {err}")
            return None

    def view_team_orders(self, team_id, check_date):
        """
        查看特定隊的排序資訊

        Args:
            team_id: 隊伍編號
            check_date: 查詢日期

        Returns:
            tuple: (是否成功, 結果資訊)
        """
        try:
            team_order = get_team_order(team_id, check_date.month)

            # 獲取隊伍成員資訊
            query = "SELECT S_ID, name, current_shift FROM Employee_Shift WHERE team = %s"
            self.db.get_cursor().execute(query, (team_id,))
            members = self.db.get_cursor().fetchall()

            if not members:
                return False, "找不到該隊資料"

            result = {
                '隊別': f'第{team_id}隊',
                '檔排序': team_order,
                '人員狀態': []
            }

            # 處理每個成員的狀態
            for member_id, name, shift_type in members:
                is_rest = not self.is_working_day(check_date, shift_type)
                shift_orders = self.get_current_shift_order(check_date)
                day_order = shift_orders.get(shift_type, 0) if not is_rest else '休假'

                result['人員狀態'].append({
                    '姓名': name,
                    '假檔': shift_type,
                    '日排序': day_order,
                    '狀態': '休假' if is_rest else '上班'
                })

            return True, result

        except Exception as err:
            return False, f"查詢錯誤: {err}"

    def check_current_rest_shift(self, check_date):
        """
        檢查指定日期哪個檔次在輪休

        Args:
            check_date: 查詢日期

        Returns:
            str: 休假狀態說明
        """
        check_date = format_date(check_date)

        # 週三特殊處理
        if check_date.weekday() == 2:
            return "今天是週三，所有檔次都在上班"

        rest_shifts = []
        all_shifts = ['123檔期', '456檔期', '789檔期']

        for shift in all_shifts:
            if not self.is_working_day(check_date, shift):
                rest_shifts.append(shift)

        if rest_shifts:
            return f"今天是 {check_date.strftime('%Y-%m-%d')}，{', '.join(rest_shifts)} 在休假"
        return "查詢出錯，請確認日期"

    def is_working_day(self, date, shift_type):
        """
        檢查指定日期是否為工作日

        Args:
            date: 查詢日期
            shift_type: 班別類型

        Returns:
            bool: 是否為工作日
        """
        date = format_date(date)
        self.shift_start_date = format_date(self.shift_start_date)

        # 週三都要上班
        if date.weekday() == 2:
            return True

        days_from_start = (date - self.shift_start_date).days
        shift_offset = self.shift_patterns[shift_type]
        adjusted_days = (days_from_start - shift_offset) % 21

        # 判斷工作週期
        if adjusted_days < 7:  # 前7天上班
            return True
        elif adjusted_days < 11:  # 接著休假4天
            return False
        elif adjusted_days < 19:  # 接著上班8天
            return True
        else:  # 最後休假2天
            return False

    def get_employee_info(self, s_id):
        """
        查詢員工資訊

        Args:
            s_id: 員工編號

        Returns:
            dict: 員工資訊
        """
        query = """
        SELECT name, job_rank, current_shift 
        FROM Employee_Shift 
        WHERE S_ID = %s
        """
        self.db.get_cursor().execute(query, (s_id,))
        result = self.db.get_cursor().fetchone()
        return {'name': result[0], 'rank': result[1], 'shift': result[2]} if result else None

    def get_employee_team(self, s_id):
        """
        獲取員工所屬的隊伍編號

        Args:
            s_id: 員工編號

        Returns:
            str: 隊伍編號
        """
        query = "SELECT team FROM Employee_Shift WHERE S_ID = %s"
        self.db.get_cursor().execute(query, (s_id,))
        result = self.db.get_cursor().fetchone()
        return result[0] if result else None

    def check_shift_assigned(self, shift_name, shift_date):
        """
        檢查該班別在指定日期是否已被分配

        Args:
            shift_name: 班別名稱
            shift_date: 日期

        Returns:
            tuple: (是否已分配, 當前分配的員工資訊)
        """
        query = """
        SELECT s.S_ID, e.name, e.team
        FROM Shift s
        JOIN Employee_Shift e ON s.S_ID = e.S_ID
        WHERE s.shift_name = %s AND s.shift_date = %s
        """
        self.db.get_cursor().execute(query, (shift_name, shift_date))
        result = self.db.get_cursor().fetchone()

        if result:
            return True, {
                'S_ID': result[0],
                'name': result[1],
                'team': result[2]
            }
        return False, None

    def modify_shift(self, shift_name, old_sid, new_sid, shift_date):
        """
        修改班別分配

        Args:
            shift_name: 班別名稱
            old_sid: 原警員編號
            new_sid: 新警員編號
            shift_date: 日期

        Returns:
            tuple: (是否成功, 結果訊息)
        """
        try:
            # 驗證新警員資訊
            new_emp_info = self.get_employee_info(new_sid)
            if not new_emp_info:
                return False, f"錯誤：找不到警員編號 {new_sid}"

            # 檢查職級限制
            rank_restrictions = get_rank_restrictions()
            if shift_name in rank_restrictions and new_emp_info['rank'] != rank_restrictions[shift_name]:
                return False, f"錯誤：{shift_name}只能由{rank_restrictions[shift_name]}擔任"

            # 檢查工作日
            shift_date = format_date(shift_date)
            if not self.is_working_day(shift_date, new_emp_info['shift']):
                return False, "錯誤：根據輪班表，該員工在此日期應該休假"

            # 檢查新警員是否已有其他班別
            check_query = """
            SELECT shift_name 
            FROM Shift 
            WHERE S_ID = %s AND shift_date = %s
            """
            self.db.get_cursor().execute(check_query, (new_sid, shift_date))
            existing_shift = self.db.get_cursor().fetchone()
            if existing_shift:
                return False, f"錯誤：該警員在此日期已被安排 {existing_shift[0]}"

            # 計算新的排序
            team = self.get_employee_team(new_sid)
            team_order = get_team_order(team, shift_date.month)
            shift_orders = self.get_current_shift_order(shift_date)
            day_order = shift_orders.get(new_emp_info['shift'], 0)

            # 更新班別
            update_query = """
            UPDATE Shift 
            SET S_ID = %s, team_order = %s, day_order = %s
            WHERE shift_name = %s AND shift_date = %s AND S_ID = %s
            """
            self.db.get_cursor().execute(update_query,
                                         (new_sid, team_order, day_order, shift_name, shift_date, old_sid))
            self.db.get_connection().commit()

            order_info = f"(檔排序: {team_order}, 日排序: {day_order})"
            return True, f"成功：已將 {shift_name} 從原警員改為 {new_emp_info['name']} {order_info}"

        except Exception as err:
            self.db.get_connection().rollback()
            return False, f"修改失敗：{str(err)}"

    def assign_shift(self, shift_name, s_id, shift_date):
        """修改後的指派班別方法"""
        try:
            print(f"\n=== 開始處理班別指派 ===")
            print(f"班別: {shift_name}")
            print(f"人員ID: {s_id}")
            print(f"日期: {shift_date}")

            # 檢查該班別是否已被分配
            is_assigned, current_emp = self.check_shift_assigned(shift_name, shift_date)
            if is_assigned:
                print(f"警告：此班別目前已由 {current_emp['name']}({current_emp['team']}隊) 擔任")
                return False, "錯誤：此班別已有人擔任，如需修改請使用修改功能"

            # 檢查該警員是否已有班別
            check_query = """
            SELECT shift_name 
            FROM Shift 
            WHERE S_ID = %s AND shift_date = %s
            """
            self.db.get_cursor().execute(check_query, (s_id, shift_date))
            existing_shift = self.db.get_cursor().fetchone()
            if existing_shift:
                return False, f"錯誤：該警員在此日期已被安排 {existing_shift[0]}"

            # [其餘驗證邏輯保持不變]
            emp_info = self.get_employee_info(s_id)
            if not emp_info:
                return False, f"錯誤：找不到警員編號 {s_id}"

            rank_restrictions = get_rank_restrictions()
            if shift_name in rank_restrictions and emp_info['rank'] != rank_restrictions[shift_name]:
                return False, f"錯誤：{shift_name}只能由{rank_restrictions[shift_name]}擔任"

            shift_date = format_date(shift_date)
            if not self.is_working_day(shift_date, emp_info['shift']):
                return False, "錯誤：根據輪班表，該員工在此日期應該休假"

            team = self.get_employee_team(s_id)
            team_order = get_team_order(team, shift_date.month)
            shift_orders = self.get_current_shift_order(shift_date)
            day_order = shift_orders.get(emp_info['shift'], 0)

            insert_query = """
            INSERT INTO Shift (shift_name, S_ID, shift_date, team_order, day_order)
            VALUES (%s, %s, %s, %s, %s)
            """
            self.db.get_cursor().execute(insert_query,
                                         (shift_name, s_id, shift_date, team_order, day_order))
            self.db.get_connection().commit()

            order_info = f"(檔排序: {team_order}, 日排序: {day_order})"
            return True, f"成功：已將 {emp_info['name']} 安排至 {shift_date} 的 {shift_name} {order_info}"

        except Exception as err:
            self.db.get_connection().rollback()
            return False, f"錯誤：{str(err)}"

    def get_current_shift_order(self, check_date):
        """
        取得當前日期各假檔的排序

        Args:
            check_date: 查詢日期

        Returns:
            dict: 各假檔的排序
        """
        return {
            shift: self.get_day_order_by_shift(shift, check_date)
            if self.is_working_day(check_date, shift) else 0
            for shift in ['123檔期', '456檔期', '789檔期']
        }

    def get_day_order_by_shift(self, shift_type, check_date):
        """
        根據假檔和日期取得排序

        Args:
            shift_type: 假檔類型
            check_date: 查詢日期

        Returns:
            int: 排序號碼
        """
        if not self.is_working_day(check_date, shift_type):
            return 0

        days_from_start = (check_date - self.shift_start_date).days
        shift_offset = self.shift_patterns[shift_type]
        adjusted_days = (days_from_start - shift_offset) % 21

        if adjusted_days < 7:
            day_patterns = [1, 2, 2, 1, 2, 1, 2]
            return day_patterns[adjusted_days]
        elif adjusted_days < 11:
            return 0
        else:
            work_day = adjusted_days - 11
            if work_day < 8:
                day_patterns = [1, 2, 1, 2, 1, 1, 2, 3]
                return day_patterns[work_day]
        return 0

    def generate_all_standby_groups(self, check_date):
        """
        生成所有可能的備勤人員分組

        Args:
            check_date: 查詢日期

        Returns:
            tuple: (是否成功, 分組結果)
        """
        try:
            shift_orders = self.get_current_shift_order(check_date)
            current_month = check_date.month

            # 取得各隊的檔排序
            team_orders = {
                str(team): get_team_order(str(team), current_month)
                for team in range(1, 15) if str(team) not in ['10', '12']
            }

            # 取得已被安排值班的人員
            duty_query = "SELECT S_ID FROM Shift WHERE shift_date = %s"
            self.db.get_cursor().execute(duty_query, (check_date,))
            duty_members = set(row[0] for row in self.db.get_cursor().fetchall())

            available_officers = []
            available_captains = []

            # 處理1-9隊的警務員
            regular_officers_query = """
            SELECT e.S_ID, e.name, e.team, e.current_shift, e.job_rank
            FROM Employee_Shift e
            WHERE e.job_rank = '警務員'
                AND e.team IN ('1','2','3','4','5','6','7','8','9')
            ORDER BY e.team
            """

            # 處理1-9隊的隊長
            regular_captains_query = """
            SELECT e.S_ID, e.name, e.team, e.current_shift, e.job_rank
            FROM Employee_Shift e
            WHERE e.job_rank = '隊長'
                AND e.team IN ('1','2','3','4','5','6','7','8','9')
            ORDER BY e.team
            """

            # 按照日排序和檔排序處理
            for day_order in [1, 2, 3]:
                for team_order in [1, 2, 3]:
                    # 處理警務員
                    self.db.get_cursor().execute(regular_officers_query)
                    for member in self.db.get_cursor().fetchall():
                        s_id, name, team, shift_type, rank = member

                        if s_id in duty_members or \
                                not self.is_working_day(check_date, shift_type) or \
                                shift_orders[shift_type] != day_order or \
                                team_orders[team] != team_order:
                            continue

                        available_officers.append({
                            'S_ID': s_id,
                            'name': name,
                            'team': team,
                            'shift_type': shift_type,
                            'team_order': team_order,
                            'day_order': day_order
                        })

                    # 處理隊長
                    self.db.get_cursor().execute(regular_captains_query)
                    for captain in self.db.get_cursor().fetchall():
                        s_id, name, team, shift_type, rank = captain

                        if s_id in duty_members or \
                                not self.is_working_day(check_date, shift_type) or \
                                shift_orders[shift_type] != day_order or \
                                team_orders[team] != team_order:
                            continue

                        available_captains.append({
                            'S_ID': s_id,
                            'name': name,
                            'team': team,
                            'shift_type': shift_type,
                            'team_order': team_order,
                            'day_order': day_order
                        })

            # 處理11,13,14隊
            special_query = """
                            SELECT e.S_ID, e.name, e.team, e.current_shift, e.job_rank
                            FROM Employee_Shift e
                            WHERE e.team IN ('11','13','14')
                            ORDER BY e.team
                            """

            self.db.get_cursor().execute(special_query)
            special_members = self.db.get_cursor().fetchall()

            for team_order in [1, 2, 3]:
                for member in special_members:
                    s_id, name, team, shift_type, rank = member

                    if s_id in duty_members or \
                            not self.is_working_day(check_date, shift_type) or \
                            team_orders[team] != team_order:
                        continue

                    member_info = {
                        'S_ID': s_id,
                        'name': name,
                        'team': team,
                        'shift_type': shift_type,
                        'team_order': team_order,
                        'day_order': 0
                    }

                    if rank == '警務員':
                        available_officers.append(member_info)
                    elif rank == '隊長':
                        available_captains.append(member_info)

            # 進行分組
            groups = []
            officer_index = 0
            group_num = 1
            last_group_captains = available_captains[:]

            # 每組9個警務員+1個隊長
            while officer_index + 9 <= len(available_officers) and len(last_group_captains) > 0:
                group = {
                    'captain': last_group_captains.pop(0),
                    'officers': available_officers[officer_index:officer_index + 9],
                    'group_num': group_num
                }
                groups.append(group)
                officer_index += 9
                group_num += 1

            # 處理剩餘人員
            if officer_index < len(available_officers) or len(last_group_captains) > 0:
                remaining_officers = available_officers[officer_index:]
                group = {
                    'captains': last_group_captains,
                    'officers': remaining_officers,
                    'group_num': group_num,
                    'is_last_group': True
                }
                groups.append(group)

            return True, groups

        except Exception as err:
            return False, f"資料庫錯誤: {str(err)}"


    def export_to_word(self, groups, check_date):
        """
        將人員列表輸出為Word文件

        Args:
            groups: 分組結果
            check_date: 日期

        Returns:
            str: 生成的檔案名稱
        """
        try:
            doc = Document()
            doc.add_heading(f'{check_date.strftime("%Y-%m-%d")} 人員列表', 0)

            # 值班人員部分
            doc.add_heading('值班人員', level=1)
            duty_query = """
                            SELECT s.shift_name, e.name, e.team
                            FROM Shift s
                            JOIN Employee_Shift e ON s.S_ID = e.S_ID
                            WHERE s.shift_date = %s
                            ORDER BY 
                                CASE shift_name
                                    WHEN 'A班' THEN '01'
                                    WHEN 'B班' THEN '02'
                                    WHEN 'C班' THEN '03'
                                    WHEN 'D班' THEN '04'
                                    WHEN 'E班' THEN '05'
                                    WHEN '上值日' THEN '06'
                                    WHEN '下值日' THEN '07'
                                    WHEN '日值日官' THEN '08'
                                    WHEN '夜值日官' THEN '09'
                                    WHEN '值班副大隊長' THEN '10'
                                    WHEN '日勤務管理員' THEN '11'
                                    WHEN '夜勤務管理員' THEN '12'
                                    WHEN '日械彈管理員' THEN '13'
                                    WHEN '夜械彈管理員' THEN '14'
                                END
                            """
            self.db.get_cursor().execute(duty_query, (check_date,))
            duty_results = self.db.get_cursor().fetchall()

            if duty_results:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                header_cells = table.rows[0].cells
                header_cells[0].text = "班別"
                header_cells[1].text = "姓名"
                header_cells[2].text = "隊別"

                for duty in duty_results:
                    row_cells = table.add_row().cells
                    row_cells[0].text = duty[0]
                    row_cells[1].text = duty[1]
                    row_cells[2].text = f"{duty[2]}隊"

            doc.add_paragraph()

            # 備勤人員部分
            doc.add_heading('備勤人員', level=1)
            for group in groups:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'

                header_cells = table.rows[0].cells
                header_cells[0].text = "帶班隊長"
                header_cells[1].text = "警務員"
                header_cells[2].text = f"備勤{group['group_num']}組"

                row_cells = table.add_row().cells

                if 'is_last_group' in group:
                    captains_text = "\n".join([
                        f"{c['name']}({c['team']}隊)"
                        for c in group['captains']
                    ])
                    row_cells[0].text = captains_text
                else:
                    captain = group['captain']
                    row_cells[0].text = f"{captain['name']}({captain['team']}隊)"

                row_cells[1].text = " ".join([
                    f"{o['name']}({o['team']}隊)"
                    for o in group['officers']
                ])
                row_cells[2].text = f"備勤{group['group_num']}組"

                doc.add_paragraph()

            filename = f"人員列表_{check_date.strftime('%Y%m%d')}.docx"
            doc.save(filename)
            return filename

        except Exception as err:
            print(f"導出文件時發生錯誤: {str(err)}")
            return None

    # 在 ShiftManager 類中新增以下方法

    def view_team_members(self, team_id):
        """
        查看指定隊伍的所有成員

        Args:
            team_id: 隊伍編號

        Returns:
            tuple: (是否成功, DataFrame或錯誤訊息)
        """
        try:
            query = """
            SELECT 
                S_ID,
                name,
                job_rank,
                current_shift,
                CASE 
                    WHEN current_shift = '123檔期' THEN '第一檔'
                    WHEN current_shift = '456檔期' THEN '第二檔'
                    WHEN current_shift = '789檔期' THEN '第三檔'
                END as shift_name
            FROM Employee_Shift
            WHERE team = %s
            ORDER BY job_rank DESC, current_shift
            """
            self.db.get_cursor().execute(query, (team_id,))
            rows = self.db.get_cursor().fetchall()

            if not rows:
                return False, f"找不到第{team_id}隊的人員資料"

            df = pd.DataFrame(rows, columns=[
                'S_ID', '姓名', '職級', '假檔', '檔次說明'
            ])
            return True, df

        except Exception as err:
            return False, f"查詢錯誤: {str(err)}"

    def update_team_member(self, s_id, new_team=None, new_shift=None):
        """
        更新隊員資料

        Args:
            s_id: 員工編號
            new_team: 新隊伍編號(可選)
            new_shift: 新假檔(可選)

        Returns:
            tuple: (是否成功, 結果訊息)
        """
        try:
            # 檢查員工是否存在
            emp_info = self.get_employee_info(s_id)
            if not emp_info:
                return False, f"找不到員工編號 {s_id}"

            # 準備更新語句
            updates = []
            values = []
            if new_team is not None:
                # 驗證新隊伍編號
                if new_team not in ['1', '2', '3', '4', '5', '6', '7', '8', '9', '11', '13', '14']:
                    return False, "無效的隊伍編號"
                updates.append("team = %s")
                values.append(new_team)

            if new_shift is not None:
                # 驗證新假檔
                if new_shift not in ['123檔期', '456檔期', '789檔期']:
                    return False, "無效的假檔"
                updates.append("current_shift = %s")
                values.append(new_shift)

            if not updates:
                return False, "沒有提供要更新的資料"

            # 建立更新查詢
            query = f"""
            UPDATE Employee_Shift 
            SET {', '.join(updates)}
            WHERE S_ID = %s
            """
            values.append(s_id)

            self.db.get_cursor().execute(query, tuple(values))
            self.db.get_connection().commit()

            # 取得更新後的資料
            self.db.get_cursor().execute(
                "SELECT name, team, current_shift FROM Employee_Shift WHERE S_ID = %s",
                (s_id,)
            )
            result = self.db.get_cursor().fetchone()

            return True, f"成功更新 {result[0]} 的資料 (隊別: {result[1]}隊, 假檔: {result[2]})"

        except Exception as err:
            self.db.get_connection().rollback()
            return False, f"更新失敗: {str(err)}"

    def bulk_update_team_shifts(self, team_id, shift_assignments):
        """
        批次更新整個隊伍的假檔

        Args:
            team_id: 隊伍編號
            shift_assignments: 字典，格式為 {s_id: new_shift}

        Returns:
            tuple: (是否成功, 結果訊息)
        """
        try:
            # 驗證所有假檔
            valid_shifts = {'123檔期', '456檔期', '789檔期'}
            for shift in shift_assignments.values():
                if shift not in valid_shifts:
                    return False, f"無效的假檔: {shift}"

            # 驗證所有員工是否屬於該隊
            s_ids = ', '.join([f"'{s_id}'" for s_id in shift_assignments.keys()])
            query = f"""
            SELECT S_ID 
            FROM Employee_Shift 
            WHERE team = %s AND S_ID IN ({s_ids})
            """
            self.db.get_cursor().execute(query, (team_id,))
            valid_members = {row[0] for row in self.db.get_cursor().fetchall()}

            invalid_members = set(shift_assignments.keys()) - valid_members
            if invalid_members:
                return False, f"以下員工不屬於第{team_id}隊: {', '.join(invalid_members)}"

            # 執行批次更新
            for s_id, new_shift in shift_assignments.items():
                query = """
                UPDATE Employee_Shift 
                SET current_shift = %s 
                WHERE S_ID = %s AND team = %s
                """
                self.db.get_cursor().execute(query, (new_shift, s_id, team_id))

            self.db.get_connection().commit()
            return True, f"成功更新第{team_id}隊 {len(shift_assignments)}位成員的假檔"

        except Exception as err:
            self.db.get_connection().rollback()
            return False, f"批次更新失敗: {str(err)}"

        # 在 ShiftManager 類中新增以下方法

    def view_team_member_order(self, team_id):
        """
        查看指定隊伍的人員順序

        Args:
            team_id: 隊伍編號

        Returns:
            tuple: (是否成功, DataFrame或錯誤訊息)
        """
        try:
            query = """
            SELECT 
                S_ID,
                name,
                job_rank,
                current_shift,
                CASE current_shift
                    WHEN '123檔期' THEN 1
                    WHEN '456檔期' THEN 2
                    WHEN '789檔期' THEN 3
                END as shift_num
            FROM Employee_Shift
            WHERE team = %s
            ORDER BY job_rank DESC, shift_num
            """
            self.db.get_cursor().execute(query, (team_id,))
            rows = self.db.get_cursor().fetchall()

            if not rows:
                return False, f"找不到第{team_id}隊的人員資料"

            # 轉換為 DataFrame 並添加順序
            data = []
            for i, row in enumerate(rows, 1):
                data.append({
                    'S_ID': row[0],
                    '順序': i,
                    '姓名': row[1],
                    '職級': row[2],
                    '假檔': row[3]
                })

            df = pd.DataFrame(data)
            return True, df

        except Exception as err:
            return False, f"查詢錯誤: {str(err)}"

    def swap_member_shifts(self, team_id, s_id1, s_id2):
        """
        交換兩個隊員的假檔

        Args:
            team_id: 隊伍編號
            s_id1: 第一個警員編號
            s_id2: 第二個警員編號

        Returns:
            tuple: (是否成功, 結果訊息)
        """
        try:
            # 檢查兩個警員是否都在同一個隊伍
            query = """
            SELECT S_ID, name, current_shift 
            FROM Employee_Shift 
            WHERE team = %s AND S_ID IN (%s, %s)
            """
            self.db.get_cursor().execute(query, (team_id, s_id1, s_id2))
            results = self.db.get_cursor().fetchall()

            if len(results) != 2:
                return False, "警員編號錯誤或不在同一個隊伍"

            # 獲取當前假檔
            shifts = {row[0]: row[2] for row in results}
            names = {row[0]: row[1] for row in results}

            # 交換假檔
            query = """
            UPDATE Employee_Shift 
            SET current_shift = CASE
                WHEN S_ID = %s THEN %s
                WHEN S_ID = %s THEN %s
            END
            WHERE S_ID IN (%s, %s) AND team = %s
            """
            self.db.get_cursor().execute(query, (
                s_id1, shifts[s_id2],
                s_id2, shifts[s_id1],
                s_id1, s_id2,
                team_id
            ))

            self.db.get_connection().commit()
            return True, f"成功交換 {names[s_id1]} 和 {names[s_id2]} 的假檔"

        except Exception as err:
            self.db.get_connection().rollback()
            return False, f"交換失敗: {str(err)}"

    def update_member_order(self, team_id, order_changes):
        """
        更新隊內人員順序

        Args:
            team_id: 隊伍編號
            order_changes: 字典，格式為 {s_id: new_order}

        Returns:
            tuple: (是否成功, 結果訊息)
        """
        try:
            # 先獲取所有隊員資訊
            query = """
            SELECT S_ID, name, job_rank, current_shift
            FROM Employee_Shift
            WHERE team = %s
            ORDER BY job_rank DESC, current_shift
            """
            self.db.get_cursor().execute(query, (team_id,))
            members = self.db.get_cursor().fetchall()

            # 創建當前成員列表
            current_members = []
            for member in members:
                current_members.append({
                    'S_ID': member[0],
                    'name': member[1],
                    'job_rank': member[2],
                    'current_shift': member[3]
                })

            # 建立新順序列表
            new_order_members = current_members.copy()

            # 根據 order_changes 排序
            for s_id, new_order in sorted(order_changes.items(), key=lambda x: x[1]):
                # 找出要移動的成員
                member = next(m for m in new_order_members if m['S_ID'] == s_id)
                # 從列表中移除
                new_order_members.remove(member)
                # 插入新位置
                new_order_members.insert(new_order - 1, member)

            # 更新資料庫
            for i, member in enumerate(new_order_members):
                if i < 3:  # 前三個人
                    new_shift = '123檔期'
                elif i < 6:  # 中間三個人
                    new_shift = '456檔期'
                else:  # 最後的人
                    new_shift = '789檔期'

                update_query = """
                UPDATE Employee_Shift 
                SET current_shift = %s 
                WHERE S_ID = %s AND team = %s
                """
                self.db.get_cursor().execute(update_query, (new_shift, member['S_ID'], team_id))

            self.db.get_connection().commit()
            return True, f"成功更新第{team_id}隊 {len(order_changes)}位成員的順序"

        except Exception as err:
            self.db.get_connection().rollback()
            return False, f"更新失敗: {str(err)}"

    def swap_member_orders(self, team_id, s_id1, s_id2):
        """
        交換兩個隊員的順序（通過交換假檔實現）

        Args:
            team_id: 隊伍編號
            s_id1: 第一個警員編號
            s_id2: 第二個警員編號

        Returns:
            tuple: (是否成功, 結果訊息)
        """
        try:
            # 檢查兩個警員是否都在同一個隊伍
            query = """
            SELECT S_ID, name, current_shift 
            FROM Employee_Shift 
            WHERE team = %s AND S_ID IN (%s, %s)
            """
            self.db.get_cursor().execute(query, (team_id, s_id1, s_id2))
            results = self.db.get_cursor().fetchall()

            if len(results) != 2:
                return False, "警員編號錯誤或不在同一個隊伍"

            # 獲取當前假檔
            shifts = {row[0]: row[2] for row in results}
            names = {row[0]: row[1] for row in results}

            # 交換假檔
            query = """
            UPDATE Employee_Shift 
            SET current_shift = CASE
                WHEN S_ID = %s THEN %s
                WHEN S_ID = %s THEN %s
            END
            WHERE S_ID IN (%s, %s) AND team = %s
            """
            self.db.get_cursor().execute(query, (
                s_id1, shifts[s_id2],
                s_id2, shifts[s_id1],
                s_id1, s_id2,
                team_id
            ))

            self.db.get_connection().commit()
            return True, f"成功交換 {names[s_id1]} 和 {names[s_id2]} 的順序"

        except Exception as err:
            self.db.get_connection().rollback()
            return False, f"交換失敗: {str(err)}"